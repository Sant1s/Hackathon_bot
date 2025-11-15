#!/usr/bin/env python3
"""
Скрипт для создания DOCX файла с ответами на вопросы о документации проекта.
"""

import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("✗ Ошибка: библиотека 'python-docx' не установлена!")
    print("  Установите её командой: pip install python-docx")
    sys.exit(1)


def add_code_paragraph(doc, text):
    """Добавляет параграф с кодом в моноширинном шрифте."""
    p = doc.add_paragraph(text, style='No Spacing')
    p.runs[0].font.name = 'Courier New'
    p.runs[0].font.size = Pt(10)
    return p


def create_requirements_document():
    """
    Создает DOCX документ с ответами на вопросы о документации проекта.
    """
    doc = Document()
    
    # Заголовок документа
    title = doc.add_heading('Документация проекта', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Раздел 1: Репозиторий GitHub
    doc.add_heading('1. Репозиторий на GitHub', level=1)
    p1 = doc.add_paragraph(
        'Ссылка на репозиторий GitHub:'
    )
    p1_link = doc.add_paragraph(
        '[Укажите ссылку на ваш репозиторий GitHub]',
        style='No Spacing'
    )
    p1_link.runs[0].font.name = 'Courier New'
    p1_link.runs[0].font.color.rgb = None  # Синий цвет для ссылки
    
    p1_desc = doc.add_paragraph(
        'В репозитории размещены все необходимые файлы проекта, включая исходный код, '
        'Docker-образ, файлы зависимостей и документацию.'
    )
    
    doc.add_paragraph()
    
    # Раздел 2: Docker-образ
    doc.add_heading('2. Docker-образ', level=1)
    p2 = doc.add_paragraph(
        'В репозитории присутствует Docker-образ (Dockerfile), который позволяет собрать '
        'и запустить приложение в контейнере.'
    )
    
    doc.add_paragraph()
    doc.add_heading('2.1. Описание Dockerfile', level=2)
    p2_1 = doc.add_paragraph(
        'Dockerfile использует многоэтапную сборку:'
    )
    dockerfile_desc = [
        'Этап сборки: использует образ golang:1.24-alpine для компиляции приложения',
        'Этап выполнения: использует образ alpine:latest с минимальными зависимостями',
        'Приложение компилируется в статический бинарный файл для Linux',
        'Порт 8080 экспонируется для доступа к приложению'
    ]
    for item in dockerfile_desc:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_paragraph()
    doc.add_heading('2.2. Расположение', level=2)
    p2_2 = doc.add_paragraph(
        'Dockerfile находится в корне репозитория и готов к использованию для сборки образа приложения.'
    )
    
    doc.add_paragraph()
    
    # Раздел 3: Requirements.txt / go.mod
    doc.add_heading('3. Описание зависимостей', level=1)
    p3 = doc.add_paragraph(
        'В корне репозитория находится файл go.mod, который содержит описание всех используемых '
        'библиотек и их версий.'
    )
    
    doc.add_paragraph()
    doc.add_heading('3.1. Основные зависимости', level=2)
    p3_1 = doc.add_paragraph(
        'Проект использует следующие основные библиотеки:'
    )
    dependencies = [
        'github.com/gorilla/mux v1.8.1 - HTTP роутер и мультиплексор',
        'github.com/lib/pq v1.10.9 - Драйвер PostgreSQL для Go',
        'github.com/minio/minio-go/v7 v7.0.66 - Клиент для работы с MinIO',
        'github.com/golang-jwt/jwt/v5 v5.3.0 - Работа с JWT токенами',
        'github.com/go-playground/validator/v10 v10.28.0 - Валидация данных',
        'github.com/joho/godotenv v1.5.1 - Загрузка переменных окружения',
        'github.com/swaggo/swag v1.16.6 - Генерация Swagger документации',
        'golang.org/x/crypto v0.44.0 - Криптографические функции'
    ]
    for dep in dependencies:
        doc.add_paragraph(dep, style='List Bullet')
    
    doc.add_paragraph()
    doc.add_heading('3.2. Версия Go', level=2)
    p3_2 = doc.add_paragraph(
        'Проект использует Go версии 1.24.0 с toolchain go1.24.3.'
    )
    
    doc.add_paragraph()
    
    # Раздел 4: README.md
    doc.add_heading('4. Файл README.md', level=1)
    p4 = doc.add_paragraph(
        'В корне репозитория находится файл README.md с подробной инструкцией по запуску проекта. '
        'README содержит следующие разделы:'
    )
    
    doc.add_paragraph()
    doc.add_heading('4.1. Описание проекта', level=2)
    p4_1 = doc.add_paragraph(
        'README содержит описание проекта, его назначение и основные возможности. '
        'Проект представляет собой HTTP сервер на Go с использованием PostgreSQL для хранения '
        'метаданных и MinIO для хранения файлов.'
    )
    
    doc.add_paragraph()
    doc.add_heading('4.2. Требования', level=2)
    p4_2 = doc.add_paragraph(
        'В README указаны системные требования:'
    )
    requirements = [
        'Go 1.21 или выше (для локального запуска)',
        'PostgreSQL (для локального запуска)',
        'MinIO (для локального запуска)',
        'Docker и Docker Compose (для запуска через контейнеры)'
    ]
    for req in requirements:
        doc.add_paragraph(req, style='List Bullet')
    
    doc.add_paragraph()
    doc.add_heading('4.3. Инструкция по локальному запуску', level=2)
    p4_3 = doc.add_paragraph(
        'README содержит подробную инструкцию по запуску проекта на локальной машине БЕЗ использования Docker:'
    )
    
    doc.add_paragraph()
    doc.add_heading('Шаг 1: Клонирование репозитория', level=3)
    step1_code = """git clone <ссылка_на_репозиторий>
cd TMPHackBackend"""
    add_code_paragraph(doc, step1_code)
    
    doc.add_paragraph()
    doc.add_heading('Шаг 2: Установка зависимостей', level=3)
    step2_code = """go mod download"""
    add_code_paragraph(doc, step2_code)
    
    doc.add_paragraph()
    doc.add_heading('Шаг 3: Настройка переменных окружения', level=3)
    step3_text = doc.add_paragraph(
        'Создайте файл .env на основе .env.example и настройте следующие переменные:'
    )
    step3_code = """cp .env.example .env
# Отредактируйте .env файл с вашими настройками"""
    add_code_paragraph(doc, step3_code)
    env_vars = [
        'PORT - порт для запуска сервера (по умолчанию 8080)',
        'DATABASE_URL - строка подключения к PostgreSQL',
        'MINIO_ENDPOINT - адрес MinIO сервера',
        'MINIO_ACCESS_KEY_ID - ключ доступа к MinIO',
        'MINIO_SECRET_ACCESS_KEY - секретный ключ MinIO',
        'JWT_SECRET - секретный ключ для JWT токенов'
    ]
    for var in env_vars:
        doc.add_paragraph(var, style='List Bullet 2')
    
    doc.add_paragraph()
    doc.add_heading('Шаг 4: Настройка базы данных', level=3)
    step4_text = doc.add_paragraph(
        'Создайте базу данных PostgreSQL и выполните миграцию (схема создается автоматически при первом запуске).'
    )
    
    doc.add_paragraph()
    doc.add_heading('Шаг 5: Запуск MinIO', level=3)
    step5_code = """docker run -p 9000:9000 -p 9001:9001 \\
  -e "MINIO_ROOT_USER=minioadmin" \\
  -e "MINIO_ROOT_PASSWORD=minioadmin" \\
  minio/minio server /data --console-address ":9001\""""
    add_code_paragraph(doc, step5_code)
    
    doc.add_paragraph()
    doc.add_heading('Шаг 6: Запуск приложения', level=3)
    step6_code = """go run ."""
    add_code_paragraph(doc, step6_code)
    
    doc.add_paragraph()
    doc.add_heading('4.4. Команды для сборки и запуска Docker-контейнера', level=2)
    p4_4 = doc.add_paragraph(
        'README содержит команды для работы с Docker-контейнерами:'
    )
    
    doc.add_paragraph()
    doc.add_heading('Сборка Docker-образа', level=3)
    build_code = """docker build -t tmphackbackend ."""
    add_code_paragraph(doc, build_code)
    
    doc.add_paragraph()
    doc.add_heading('Запуск контейнера', level=3)
    run_code = """docker run -d -p 8080:8080 --env-file .env tmphackbackend"""
    add_code_paragraph(doc, run_code)
    
    doc.add_paragraph()
    doc.add_heading('Запуск через Docker Compose', level=3)
    p4_4_1 = doc.add_paragraph(
        'Для запуска всех сервисов (приложение, PostgreSQL, MinIO) используется docker-compose:'
    )
    compose_code = """docker-compose up -d"""
    add_code_paragraph(doc, compose_code)
    
    doc.add_paragraph()
    doc.add_heading('Просмотр логов', level=3)
    logs_code = """docker-compose logs -f"""
    add_code_paragraph(doc, logs_code)
    
    doc.add_paragraph()
    doc.add_heading('Остановка контейнеров', level=3)
    stop_code = """docker-compose down"""
    add_code_paragraph(doc, stop_code)
    
    doc.add_paragraph()
    doc.add_heading('4.5. Пример запуска через командную строку', level=2)
    p4_5 = doc.add_paragraph(
        'В README приведены готовые примеры команд для различных сценариев запуска:'
    )
    
    doc.add_paragraph()
    doc.add_heading('Пример 1: Полный локальный запуск', level=3)
    example1 = """# 1. Клонирование репозитория
git clone https://github.com/your-username/TMPHackBackend.git
cd TMPHackBackend

# 2. Установка зависимостей
go mod download

# 3. Настройка переменных окружения
cp .env.example .env
# Отредактируйте .env файл

# 4. Запуск MinIO (в отдельном терминале)
docker run -d -p 9000:9000 -p 9001:9001 \\
  -e "MINIO_ROOT_USER=minioadmin" \\
  -e "MINIO_ROOT_PASSWORD=minioadmin" \\
  minio/minio server /data --console-address ":9001"

# 5. Запуск приложения
go run ."""
    add_code_paragraph(doc, example1)
    
    doc.add_paragraph()
    doc.add_heading('Пример 2: Запуск через Docker', level=3)
    example2 = """# 1. Клонирование репозитория
git clone https://github.com/your-username/TMPHackBackend.git
cd TMPHackBackend

# 2. Настройка переменных окружения
cp .env.example .env
# Отредактируйте .env файл

# 3. Сборка образа
docker build -t tmphackbackend .

# 4. Запуск контейнера
docker run -d -p 8080:8080 --env-file .env tmphackbackend"""
    add_code_paragraph(doc, example2)
    
    doc.add_paragraph()
    doc.add_heading('Пример 3: Запуск через Docker Compose', level=3)
    example3 = """# 1. Клонирование репозитория
git clone https://github.com/your-username/TMPHackBackend.git
cd TMPHackBackend

# 2. Настройка переменных окружения
cp .env.example .env
# Отредактируйте .env файл

# 3. Запуск всех сервисов
docker-compose up -d

# 4. Проверка статуса
docker-compose ps

# 5. Просмотр логов
docker-compose logs -f

# 6. Остановка всех сервисов
docker-compose down"""
    add_code_paragraph(doc, example3)
    
    doc.add_paragraph()
    doc.add_heading('4.6. Дополнительная информация в README', level=2)
    p4_6 = doc.add_paragraph(
        'README также содержит:'
    )
    additional = [
        'Описание API endpoints (Swagger документация доступна по адресу /swagger/)',
        'Структура проекта',
        'Информация о лицензии',
        'Health check endpoint для проверки состояния сервера'
    ]
    for item in additional:
        doc.add_paragraph(item, style='List Bullet')
    
    # Сохраняем документ
    output_path = Path(__file__).parent.parent / "Требования_к_документации.docx"
    doc.save(str(output_path))
    
    print(f"✓ Документ создан: {output_path}")
    return output_path


if __name__ == "__main__":
    create_requirements_document()
